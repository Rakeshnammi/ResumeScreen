import PyPDF2
import docx
import os
import re
from typing import Dict, Optional, List
from data_extractor import DataExtractor

class ResumeProcessor:
    """Handles resume file processing and text extraction"""
    
    def __init__(self):
        self.data_extractor = DataExtractor()
    
    def extract_resume_data(self, file_path: str, filename: str) -> Optional[Dict]:
        """
        Extract data from resume file with improved preprocessing
        
        Args:
            file_path: Path to the resume file
            filename: Original filename
            
        Returns:
            Dictionary containing extracted resume data
        """
        try:
            # Extract text based on file type
            raw_text = self._extract_text_from_file(file_path)
            
            if not raw_text or len(raw_text.strip()) < 50:
                return None
            
            # Clean and preprocess text for better extraction
            cleaned_text = self._preprocess_text(raw_text)
            
            # Validate that it looks like a resume
            if not self.validate_resume_content(cleaned_text):
                print(f"Warning: {filename} may not be a resume")
            
            # Extract structured data using NLP
            extracted_data = self.data_extractor.extract_candidate_data(cleaned_text)
            
            # Add metadata
            extracted_data.update({
                'filename': filename,
                'raw_text': cleaned_text,
                'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0
            })
            
            return extracted_data
            
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")
            return None
    
    def _preprocess_text(self, text: str) -> str:
        """
        Preprocess extracted text for better data extraction
        
        Args:
            text: Raw extracted text
            
        Returns:
            Preprocessed text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace while preserving line breaks
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove page numbers and common footer/header artifacts
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Normalize common resume section headers
        section_headers = {
            r'work\s+experience': 'EXPERIENCE',
            r'professional\s+experience': 'EXPERIENCE',
            r'employment\s+history': 'EXPERIENCE',
            r'education': 'EDUCATION',
            r'academic\s+background': 'EDUCATION',
            r'skills': 'SKILLS',
            r'technical\s+skills': 'SKILLS',
            r'core\s+competencies': 'SKILLS',
            r'certifications?': 'CERTIFICATIONS',
            r'projects?': 'PROJECTS',
        }
        
        for pattern, replacement in section_headers.items():
            text = re.sub(f'(?i)^{pattern}\\s*:?\\s*$', replacement, text, flags=re.MULTILINE)
        
        # Remove special characters that may interfere with parsing
        text = text.replace('\x00', '')  # Remove null characters
        text = text.replace('\uf0b7', '•')  # Normalize bullet points
        
        return text.strip()
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from PDF or Word document
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text content
        """
        text = ""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                text = self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            
        return text
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for better accuracy"""
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text with layout preservation
                    page_text = page.extract_text(layout=True)
                    if page_text:
                        text += page_text + "\n"
                    
                    # Also extract tables if present
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " ".join([str(cell) for cell in row if cell]) + "\n"
                
                if text.strip():  # If pdfplumber succeeded, return
                    return text
        except Exception as e:
            print(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2...")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
        except Exception as e:
            print(f"Error reading PDF {file_path}: {str(e)}")
        
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document with improved formatting"""
        text = ""
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs with formatting preservation
            for paragraph in doc.paragraphs:
                # Preserve paragraph structure
                para_text = paragraph.text.strip()
                if para_text:
                    # Add bullet points if it's a list item
                    if paragraph.style.name.startswith('List'):
                        text += "• " + para_text + "\n"
                    else:
                        text += para_text + "\n"
            
            # Extract text from tables with better structure
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"  # Add spacing after table
            
            # Extract headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        header_text = paragraph.text.strip()
                        if header_text:
                            text += header_text + "\n"
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        footer_text = paragraph.text.strip()
                        if footer_text:
                            text += footer_text + "\n"
                    
        except Exception as e:
            print(f"Error reading Word document {file_path}: {str(e)}")
        
        return text
    
    def validate_resume_content(self, text: str) -> bool:
        """
        Validate if the extracted text appears to be a resume
        
        Args:
            text: Extracted text content
            
        Returns:
            True if content appears to be a resume
        """
        if not text or len(text.strip()) < 100:
            return False
        
        # Check for common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'job', 'position',
            'resume', 'cv', 'curriculum vitae', 'qualifications'
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        # Should have at least 3 resume indicators
        return indicator_count >= 3
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\@\(\)\[\]\/]', '', text)
        
        # Remove multiple consecutive dots or dashes
        text = re.sub(r'\.{3,}', '...', text)
        text = re.sub(r'-{3,}', '---', text)
        
        return text.strip()
